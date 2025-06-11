"""
文档处理模块

负责对已生成的文档进行二次处理，如删除文件名、删除注释等。
"""

import os
import re
import docx
from typing import Dict, List, Tuple
from docx import Document
import datetime

from .code_processor import is_english_text, is_end_of_line_comment


def is_filename_line(line: str) -> bool:
    """
    判断一行是否为文件名行
    
    Args:
        line: 文档中的一行文本
        
    Returns:
        是否为文件名行
    """
    # 文件名通常是单独的一行，不包含特殊格式
    # 可以根据实际文件名格式进行调整
    if not line.strip():
        return False
    
    # 检查是否是常见的文件扩展名结尾
    common_extensions = ['.py', '.java', '.c', '.cpp', '.cs', '.js', '.html', '.css', '.php', '.go', '.rb', '.swift']
    for ext in common_extensions:
        if line.strip().endswith(ext):
            return True
    
    return False


def is_comment_line(line: str) -> bool:
    """
    判断一行是否为注释行
    
    Args:
        line: 文档中的一行文本
        
    Returns:
        是否为注释行
    """
    stripped = line.strip()
    
    # 检查是否为空行
    if not stripped:
        return False
    
    # 检查是否为Python注释
    if stripped.startswith('#'):
        return True
    
    # 检查是否为C/C++/Java/JavaScript等语言的单行注释
    if stripped.startswith('//'):
        return True
    
    # 检查是否为Python三引号注释（开始或结束）
    if stripped.startswith('"""') or stripped.startswith("'''"):
        return True
    if stripped.endswith('"""') or stripped.endswith("'''"):
        return True
    
    # 检查是否为HTML/XML注释
    if stripped.startswith('<!--') and stripped.endswith('-->'):
        return True
    
    # 检查是否为整行的C/C++/Java多行注释
    if (stripped.startswith('/*') and stripped.endswith('*/')) or \
       (stripped.startswith('/**') and stripped.endswith('*/')):
        return True
    
    # 检查是否为多行注释的一部分
    if stripped.startswith('*') and not stripped.startswith('*/'):
        return True
    
    return False


def process_document(
    input_path: str, 
    output_path: str, 
    deleted_content_path: str, 
    remove_filenames: bool = False, 
    remove_large_comments: bool = False, 
    remove_english_comments: bool = False,
    remove_comments_ratio: int = 0
) -> Tuple[str, str, Dict]:
    """
    处理已生成的文档，根据选择条件删除文件名和注释
    
    Args:
        input_path: 输入文档路径
        output_path: 输出文档路径
        deleted_content_path: 删除内容记录文档路径
        remove_filenames: 是否删除文件名
        remove_large_comments: 是否删除大段注释
        remove_english_comments: 是否删除英文注释
        remove_comments_ratio: 随机删除注释的比例，0表示不删除
        
    Returns:
        元组(输出文档路径, 删除内容记录文档路径, 统计信息)
    """
    # 检查输入文件是否存在
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"输入文件不存在: {input_path}")
    
    # 读取原始文档
    doc = Document(input_path)
    
    # 创建新文档
    new_doc = Document()
    deleted_doc = Document()
    
    # 添加删除内容记录文档的标题
    deleted_doc.add_heading('删除内容记录', 0)
    deleted_doc.add_paragraph(f"原始文档: {os.path.basename(input_path)}")
    deleted_doc.add_paragraph(f"删除选项:")
    deleted_doc.add_paragraph(f"- 删除文件名: {'是' if remove_filenames else '否'}")
    deleted_doc.add_paragraph(f"- 删除大段注释: {'是' if remove_large_comments else '否'}")
    deleted_doc.add_paragraph(f"- 删除英文注释: {'是' if remove_english_comments else '否'}")
    deleted_doc.add_paragraph(f"- 随机删除注释比例: {remove_comments_ratio}")
    deleted_doc.add_paragraph("")
    deleted_doc.add_heading('被删除的内容:', 1)
    
    # 统计信息
    stats = {
        'total_lines': 0,
        'deleted_filenames': 0,
        'deleted_large_comments': 0,
        'deleted_english_comments': 0,
        'deleted_random_comments': 0,
        'remaining_lines': 0
    }
    
    # 用于跟踪大段注释
    in_large_comment = False
    large_comment_lines = []
    large_comment_start_index = -1
    
    # 标记是否有删除内容
    has_deleted_content = False
    
    # 处理文档的每一段落
    for i, para in enumerate(doc.paragraphs):
        line = para.text
        stats['total_lines'] += 1
        
        # 跳过空行
        if not line.strip():
            new_doc.add_paragraph(line)
            continue
        
        # 处理文件名行
        if remove_filenames and is_filename_line(line):
            deleted_doc.add_paragraph(f"文件名 (行 {i+1}): {line}")
            stats['deleted_filenames'] += 1
            has_deleted_content = True
            continue
        
        # 处理注释行
        if is_comment_line(line):
            # 处理大段注释
            if remove_large_comments:
                if not in_large_comment:
                    # 可能是大段注释的开始
                    in_large_comment = True
                    large_comment_lines = [line]
                    large_comment_start_index = i
                else:
                    # 已经在大段注释中
                    large_comment_lines.append(line)
                
                # 检查下一行是否还是注释
                next_is_comment = False
                if i + 1 < len(doc.paragraphs):
                    next_line = doc.paragraphs[i+1].text
                    next_is_comment = is_comment_line(next_line)
                
                # 如果下一行不是注释，则当前大段注释结束
                if not next_is_comment:
                    in_large_comment = False
                    
                    # 如果大段注释超过1行，则删除
                    if len(large_comment_lines) > 1:
                        deleted_doc.add_heading(f"大段注释 (行 {large_comment_start_index+1}-{i+1}):", 2)
                        for comment_line in large_comment_lines:
                            deleted_doc.add_paragraph(comment_line)
                        stats['deleted_large_comments'] += len(large_comment_lines)
                        has_deleted_content = True
                        continue
                    else:
                        # 单行注释，根据其他条件决定是否保留
                        pass
            
            # 处理英文注释
            if remove_english_comments and not is_end_of_line_comment(line):
                # 提取注释内容
                comment_text = ""
                stripped = line.strip()
                
                if stripped.startswith('#'):
                    comment_text = stripped[1:].strip()
                elif stripped.startswith('//'):
                    comment_text = stripped[2:].strip()
                elif stripped.startswith('/*') and stripped.endswith('*/'):
                    comment_text = stripped[2:-2].strip()
                elif stripped.startswith('"""') and stripped.endswith('"""'):
                    comment_text = stripped[3:-3].strip()
                elif stripped.startswith("'''") and stripped.endswith("'''"):
                    comment_text = stripped[3:-3].strip()
                
                # 如果是英文注释，删除
                if comment_text and is_english_text(comment_text):
                    deleted_doc.add_paragraph(f"英文注释 (行 {i+1}): {line}")
                    stats['deleted_english_comments'] += 1
                    has_deleted_content = True
                    continue
        
        # 如果代码行到达这里，表示需要保留
        new_doc.add_paragraph(line)
        stats['remaining_lines'] += 1
    
    # 如果没有删除任何内容，添加一条消息
    if not has_deleted_content:
        deleted_doc.add_paragraph("未删除任何内容。")
    
    # 保存处理后的文档
    try:
        # 使用绝对路径
        abs_output_path = os.path.abspath(output_path)
        abs_deleted_content_path = os.path.abspath(deleted_content_path)
        
        print(f"处理文档任务完成，准备保存文件")
        print(f"输出目录权限检查:")
        output_dir = os.path.dirname(abs_output_path)
        print(f"  输出目录: {output_dir}")
        print(f"  目录存在: {os.path.exists(output_dir)}")
        
        if not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
            print(f"  已创建输出目录: {output_dir}")
            
        print(f"  目录可写: {os.access(output_dir, os.W_OK)}")
        
        # 先保存处理后的文档
        print(f"正在保存处理后文档，绝对路径: {abs_output_path}")
        new_doc.save(abs_output_path)
        print(f"处理后文档已保存，文件大小: {os.path.getsize(abs_output_path)} 字节")
        
        # 准备保存删除内容文档
        print(f"正在准备保存删除内容记录文档")
        # 强制确保删除内容文档存在，即使没有删除内容
        deleted_doc.add_paragraph(f"处理时间: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        deleted_doc.add_paragraph(f"对应的处理后文档: {os.path.basename(abs_output_path)}")
        
        # 保存前再次检查文档内容
        print(f"删除内容文档段落数: {len(deleted_doc.paragraphs)}")
        for i, para in enumerate(deleted_doc.paragraphs[:5]):  # 只打印前5个段落
            print(f"  段落 {i+1}: {para.text[:50]}..." if len(para.text) > 50 else f"  段落 {i+1}: {para.text}")
            
        # 保存删除内容文档 - 使用多种尝试方法
        saved_deleted_doc = False
        
        # 方法1: 使用原始路径
        try:
            print(f"方法1: 尝试使用原始路径保存删除内容文档: {deleted_content_path}")
            deleted_doc.save(deleted_content_path)
            print(f"方法1成功: 删除内容记录文档已保存")
            saved_deleted_doc = True
            actual_deleted_path = deleted_content_path
        except Exception as e1:
            print(f"方法1失败: {str(e1)}")
            
            # 方法2: 使用绝对路径
            try:
                print(f"方法2: 尝试使用绝对路径保存: {abs_deleted_content_path}")
                deleted_doc.save(abs_deleted_content_path)
                print(f"方法2成功: 删除内容记录文档已保存")
                saved_deleted_doc = True
                actual_deleted_path = abs_deleted_content_path
            except Exception as e2:
                print(f"方法2失败: {str(e2)}")
                
                # 方法3: 修改文件名后尝试
                try:
                    alt_filename = f"alt_{os.path.basename(deleted_content_path)}"
                    alt_path = os.path.join(output_dir, alt_filename)
                    print(f"方法3: 尝试使用替代文件名: {alt_path}")
                    deleted_doc.save(alt_path)
                    print(f"方法3成功: 删除内容记录文档已保存为替代文件")
                    saved_deleted_doc = True
                    actual_deleted_path = alt_path
                except Exception as e3:
                    print(f"方法3失败: {str(e3)}")
                    
                    # 方法4: 保存到临时目录
                    try:
                        import tempfile
                        temp_dir = tempfile.gettempdir()
                        temp_file = os.path.join(temp_dir, f"temp_{os.path.basename(deleted_content_path)}")
                        print(f"方法4: 尝试保存到临时目录: {temp_file}")
                        deleted_doc.save(temp_file)
                        print(f"方法4成功: 删除内容记录文档已保存到临时目录")
                        
                        # 尝试复制回原始位置
                        try:
                            import shutil
                            shutil.copy2(temp_file, deleted_content_path)
                            print(f"已成功将临时文件复制到目标位置: {deleted_content_path}")
                            saved_deleted_doc = True
                            actual_deleted_path = deleted_content_path
                        except Exception as e4_copy:
                            print(f"无法复制临时文件到目标位置: {str(e4_copy)}")
                            saved_deleted_doc = True
                            actual_deleted_path = temp_file
                    except Exception as e4:
                        print(f"方法4失败: {str(e4)}")
        
        # 检查是否成功保存了删除内容文档
        if saved_deleted_doc:
            print(f"删除内容记录文档已成功保存为: {actual_deleted_path}")
            if os.path.exists(actual_deleted_path):
                file_size = os.path.getsize(actual_deleted_path)
                print(f"文件大小: {file_size} 字节")
            
            # 返回处理后的文档路径和实际保存的删除内容文档路径
            return abs_output_path, actual_deleted_path, stats
        else:
            print(f"警告：所有尝试都失败，无法保存删除内容记录文档")
            # 创建一个内存中的错误报告
            error_info = f"无法保存删除内容记录文档到 {deleted_content_path}，所有尝试方法都失败。"
            # 仍然返回路径，但用户需要知道文件可能不存在
            return abs_output_path, deleted_content_path, stats
            
    except Exception as e:
        print(f"保存文档时出错: {str(e)}")
        import traceback
        traceback.print_exc()
        raise
    
    return output_path, deleted_content_path, stats


def process_document_batch(
    input_dir: str,
    output_dir: str,
    remove_filenames: bool = False,
    remove_large_comments: bool = False,
    remove_english_comments: bool = False,
    remove_comments_ratio: int = 0
) -> List[Dict]:
    """
    批量处理文档文件夹
    
    Args:
        input_dir: 输入文档文件夹
        output_dir: 输出文档文件夹
        remove_filenames: 是否删除文件名
        remove_large_comments: 是否删除大段注释
        remove_english_comments: 是否删除英文注释
        remove_comments_ratio: 随机删除注释的比例，0表示不删除
        
    Returns:
        处理结果列表
    """
    # 确保输出目录存在
    os.makedirs(output_dir, exist_ok=True)
    
    results = []
    
    # 遍历输入目录中的所有docx文件
    for filename in os.listdir(input_dir):
        if filename.endswith('.docx'):
            input_path = os.path.join(input_dir, filename)
            
            # 生成输出文件名
            base_name = os.path.splitext(filename)[0]
            output_filename = f"{base_name}_processed.docx"
            deleted_content_filename = f"{base_name}_deleted_content.docx"
            
            output_path = os.path.join(output_dir, output_filename)
            deleted_content_path = os.path.join(output_dir, deleted_content_filename)
            
            try:
                # 处理文档
                _, _, stats = process_document(
                    input_path,
                    output_path,
                    deleted_content_path,
                    remove_filenames,
                    remove_large_comments,
                    remove_english_comments,
                    remove_comments_ratio
                )
                
                # 记录结果
                results.append({
                    'input_file': input_path,
                    'output_file': output_path,
                    'deleted_content_file': deleted_content_path,
                    'stats': stats,
                    'success': True
                })
            except Exception as e:
                results.append({
                    'input_file': input_path,
                    'error': str(e),
                    'success': False
                })
    
    return results 