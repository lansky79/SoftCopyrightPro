"""
文档生成模块

负责生成Word和PDF格式的代码文档，设置等宽字体、字号和页面布局。
"""

import os
from typing import List, Optional, Dict, Tuple
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml
from docx.shared import Mm, Twips


def set_paragraph_spacing(paragraph, before=0, after=2.3, line=10.5):
    """
    设置段落间距
    
    Args:
        paragraph: 段落对象
        before: 段前间距（磅）
        after: 段后间距（磅）
        line: 行间距（磅）
    """
    p_format = paragraph.paragraph_format
    p_format.space_before = Pt(before)
    p_format.space_after = Pt(after)
    p_format.line_spacing = Pt(line)


def add_page_number(run):
    """
    添加页码字段到run
    
    Args:
        run: 文本run对象
    """
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def add_header_border(paragraph):
    """
    为段落添加下边框线
    
    Args:
        paragraph: 段落对象
    """
    # 创建下边框
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')  # 恢复原始边框宽度
    bottom.set(qn('w:space'), '0')  # 恢复原始间距，使其紧贴文本
    bottom.set(qn('w:color'), '000000')  # 边框颜色
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_chinese_font(run, font_name_cn):
    """
    强制设置中文字体
    
    Args:
        run: 文本run对象
        font_name_cn: 中文字体名称
    """
    # 设置东亚文字字体（中文）
    run.font.east_asia = font_name_cn
    
    # 在XML级别设置字体
    r = run._r
    rPr = r.get_or_add_rPr()
    
    # 设置东亚文字字体（中文）
    eastAsia = OxmlElement('w:eastAsia')
    eastAsia.set(qn('w:val'), font_name_cn)
    
    # 删除任何现有的eastAsia元素
    for element in rPr.findall(qn('w:eastAsia')):
        rPr.remove(element)
    
    rPr.append(eastAsia)


def create_aligned_header(section, software_name, software_version, font_name_en, font_name_cn):
    """
    创建对齐的页眉，确保左右两侧高度一致，分割线紧贴页眉
    
    Args:
        section: 文档节
        software_name: 软件名称
        software_version: 软件版本号
        font_name_en: 英文字体名称
        font_name_cn: 中文字体名称
    
    Returns:
        页眉对象
    """
    header = section.header
    
    # 清除现有内容
    for p in header.paragraphs:
        p._element.getparent().remove(p._element)
        p._p = None
        p._element = None
    
    # 添加空段落，使页眉下移
    empty_para = header.add_paragraph()
    empty_para.paragraph_format.space_before = Pt(0)
    empty_para.paragraph_format.space_after = Pt(0.8)  # 从1.0pt进一步减少到0.8pt，使页眉更上移
    
    # 创建单个段落用于页眉
    header_para = header.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_para.paragraph_format.space_before = Pt(0)
    header_para.paragraph_format.space_after = Pt(0)
    
    # 添加软件名称和版本号（左侧）- 移除开发者信息
    left_run = header_para.add_run(f"{software_name} {software_version}")
    left_run.font.size = Pt(9)
    left_run.font.name = font_name_en
    set_chinese_font(left_run, font_name_cn)
    
    # 添加制表位，用于右对齐页码
    tab_stops = header_para.paragraph_format.tab_stops
    tab_stop = tab_stops.add_tab_stop(
        Inches(section.page_width.inches - section.left_margin.inches - section.right_margin.inches),
        WD_ALIGN_PARAGRAPH.RIGHT
    )
    
    # 添加制表符和页码（右侧）
    tab_run = header_para.add_run("\t")
    page_run = header_para.add_run()
    add_page_number(page_run)
    page_run.font.size = Pt(9)
    page_run.font.name = font_name_en
    set_chinese_font(page_run, font_name_cn)
    
    # 为页眉段落添加下边框线
    add_header_border(header_para)
    
    return header


def create_word_document(
    content: List[str],
    output_path: str,
    file_list: List[str] = None,
    software_name: str = "SoftCopyrightPro",
    software_version: str = "v0.1.0",
    font_name_en: str = 'Courier New',
    font_name_cn: str = 'SimSun',  # 宋体
    font_size: float = 10.5,  # 调整为10.5pt
    page_width: float = 8.5,
    page_height: float = 11.0,
    margin: float = 1.0
) -> str:
    """
    创建Word文档
    
    Args:
        content: 代码内容行列表
        output_path: 输出文件路径
        file_list: 文件列表，按处理顺序排列
        software_name: 软件名称（用于页眉）
        software_version: 软件版本号（用于页眉）
        font_name_en: 英文等宽字体名称
        font_name_cn: 中文等宽字体名称
        font_size: 字体大小（磅）
        page_width: 页面宽度（英寸）
        page_height: 页面高度（英寸）
        margin: 页面边距（英寸）
    
    Returns:
        生成的文件路径
    """
    # 创建Word文档
    doc = Document()
    
    # 设置文档默认字体
    doc.styles['Normal'].font.name = font_name_en
    doc.styles['Normal'].font.size = Pt(font_size)
    
    # 在XML级别设置默认字体
    element = doc.styles['Normal']._element
    rPr = element.get_or_add_rPr()
    
    # 设置默认字体
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name_en)
    rFonts.set(qn('w:hAnsi'), font_name_en)
    rFonts.set(qn('w:eastAsia'), font_name_cn)
    rFonts.set(qn('w:cs'), font_name_en)
    
    # 删除任何现有的rFonts元素
    for old_rFonts in rPr.findall(qn('w:rFonts')):
        rPr.remove(old_rFonts)
    
    rPr.append(rFonts)
    
    # 设置页面大小和边距
    section = doc.sections[0]
    section.page_width = Inches(page_width)
    section.page_height = Inches(page_height)
    section.left_margin = Inches(margin)
    section.right_margin = Inches(margin)
    section.top_margin = Inches(margin)  # 恢复为原始值1.0英寸
    section.bottom_margin = Inches(margin)
    
    # 创建对齐的页眉，确保左右两侧高度一致，分割线紧贴页眉
    create_aligned_header(section, software_name, software_version, font_name_en, font_name_cn)
    
    # 添加内容
    for line in content:
        paragraph = doc.add_paragraph(line)
        paragraph.style.font.name = font_name_en
        paragraph.style.font.size = Pt(font_size)
        
        # 设置段落间距 - 段前0pt，段后2.3pt，行距10.5pt
        set_paragraph_spacing(paragraph, before=0, after=2.3, line=10.5)
        
        # 设置中文字体
        if paragraph.runs:
            run = paragraph.runs[0]
            run.font.name = font_name_en  # 英文字体
            set_chinese_font(run, font_name_cn)  # 中文字体
            run.font.size = Pt(font_size)
        else:
            # 如果runs为空，添加一个run
            run = paragraph.add_run(line)
            run.font.name = font_name_en  # 英文字体
            set_chinese_font(run, font_name_cn)  # 中文字体
            run.font.size = Pt(font_size)
            # 清除段落原有文本，避免重复
            paragraph.text = ""
    
    # 确保输出目录存在
    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    
    # 保存文档
    doc.save(output_path)
    
    return output_path


def convert_word_to_pdf(word_path: str, pdf_path: Optional[str] = None) -> str:
    """
    将Word文档转换为PDF
    
    Args:
        word_path: Word文档路径
        pdf_path: PDF输出路径，如果为None，则使用与Word文档相同的名称但扩展名为.pdf
    
    Returns:
        生成的PDF文件路径
    """
    if pdf_path is None:
        pdf_path = os.path.splitext(word_path)[0] + '.pdf'
    
    try:
        # 尝试使用docx2pdf
        from docx2pdf import convert
        convert(word_path, pdf_path)
        return pdf_path
    except ImportError:
        try:
            # 尝试使用win32com (仅适用于Windows)
            import win32com.client
            word = win32com.client.Dispatch('Word.Application')
            doc = word.Documents.Open(word_path)
            doc.SaveAs(pdf_path, FileFormat=17)  # FileFormat=17 表示PDF
            doc.Close()
            word.Quit()
            return pdf_path
        except ImportError:
            print("无法将Word转换为PDF。请安装docx2pdf或确保在Windows系统上运行。")
            return word_path


def generate_document(
    content: List[str],
    output_path: str,
    file_list: List[str] = None,
    output_format: str = 'docx',
    software_name: str = "SoftCopyrightPro",
    software_version: str = "v0.1.0",
    font_name_en: str = 'Courier New',
    font_name_cn: str = 'SimSun',  # 宋体
    font_size: float = 10.5  # 调整为10.5pt
) -> Tuple[str, Dict]:
    """
    生成文档
    
    Args:
        content: 代码内容行列表
        output_path: 输出文件路径
        file_list: 文件列表，按处理顺序排列
        output_format: 输出格式，'docx'或'pdf'
        software_name: 软件名称（用于页眉）
        software_version: 软件版本号（用于页眉）
        font_name_en: 英文等宽字体名称
        font_name_cn: 中文等宽字体名称
        font_size: 字体大小（磅）
    
    Returns:
        元组(生成的文件路径, 文档统计信息)
    """
    # 确保输出路径有正确的扩展名，但保留完整的文件名
    # 如果已经有扩展名，直接使用；否则添加扩展名
    if not output_path.lower().endswith('.docx'):
        word_path = output_path + '.docx'
    else:
        word_path = output_path
    
    # 创建Word文档
    word_path = create_word_document(
        content=content,
        output_path=word_path,
        file_list=file_list,
        software_name=software_name,
        software_version=software_version,
        font_name_en=font_name_en,
        font_name_cn=font_name_cn,
        font_size=font_size
    )
    
    # 计算统计信息
    stats = {
        'total_lines': len(content),
        'total_pages': (len(content) + 49) // 50,  # 估算页数，每页约50行
        'file_list': file_list or []
    }
    
    # 如果需要PDF格式，则转换
    if output_format.lower() == 'pdf':
        # 将.docx替换为.pdf，如果没有.docx扩展名，直接添加.pdf
        if word_path.lower().endswith('.docx'):
            pdf_path = word_path[:-5] + '.pdf'
        else:
            pdf_path = word_path + '.pdf'
        return convert_word_to_pdf(word_path, pdf_path), stats
    
    return word_path, stats 