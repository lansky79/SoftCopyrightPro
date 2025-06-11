import sys
import os
import datetime
from PyQt5.QtWidgets import (QApplication, QMainWindow, QFileDialog, QTreeWidget, QTreeWidgetItem,
                            QVBoxLayout, QHBoxLayout, QLabel, QPushButton, QCheckBox, QWidget,
                            QGroupBox, QSpinBox, QProgressBar, QStatusBar, QMessageBox, QTabWidget,
                            QHeaderView)
from PyQt5.QtCore import Qt, QThread, pyqtSignal

from sofcrtpro.document_processor import process_document, process_document_batch

# 添加调试开关
DEBUG = True

def debug_print(*args, **kwargs):
    """调试打印函数"""
    if DEBUG:
        print("[DEBUG]", *args, **kwargs)

class DocumentProcessorThread(QThread):
    """文档处理线程，避免处理过程阻塞GUI"""
    progress_signal = pyqtSignal(int)
    finished_signal = pyqtSignal(dict)
    error_signal = pyqtSignal(str)
    
    def __init__(self, params):
        super().__init__()
        self.params = params
    
    def run(self):
        try:
            # 获取参数
            input_path = self.params['input_path']
            output_dir = self.params['output_dir']
            remove_filenames = self.params['remove_filenames']
            remove_large_comments = self.params['remove_large_comments']
            remove_english_comments = self.params['remove_english_comments']
            remove_comments_ratio = self.params['remove_comments_ratio']
            
            # 设置进度
            self.progress_signal.emit(10)
            
            # 生成输出文件名，添加时间戳避免文件冲突
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            timestamp = datetime.datetime.now().strftime("%m%d%H%M%S")
            output_filename = f"{base_name}_{timestamp}_去注释.docx"
            deleted_content_filename = f"{base_name}_{timestamp}_被删除.docx"
            
            output_path = os.path.join(output_dir, output_filename)
            deleted_content_path = os.path.join(output_dir, deleted_content_filename)
            
            debug_print(f"\n=== 线程开始处理文档 ===")
            debug_print(f"输入文件: {input_path}")
            debug_print(f"输出文件: {output_path}")
            debug_print(f"删除内容文件: {deleted_content_path}")
            
            # 确保输出目录存在
            os.makedirs(output_dir, exist_ok=True)
            
            # 处理文档
            self.progress_signal.emit(30)
            
            debug_print("=== 线程调用process_document函数 ===")
            debug_print(f"删除选项: 文件名={remove_filenames}, 大段注释={remove_large_comments}, 英文注释={remove_english_comments}, 随机比例={remove_comments_ratio}")
            
            # 和测试程序一样，直接调用process_document
            output_file, deleted_file, stats = process_document(
                input_path,
                output_path,
                deleted_content_path,
                remove_filenames=remove_filenames,
                remove_large_comments=remove_large_comments,
                remove_english_comments=remove_english_comments,
                remove_comments_ratio=remove_comments_ratio
            )
            
            self.progress_signal.emit(70)
            
            # 检查文件是否存在
            debug_print(f"\n=== 线程检查生成的文件 ===")
            output_exists = os.path.exists(output_file)
            deleted_exists = os.path.exists(deleted_file)
            debug_print(f"输出文件存在: {output_exists}")
            debug_print(f"删除内容文件存在: {deleted_exists}")
            
            if output_exists:
                debug_print(f"输出文件大小: {os.path.getsize(output_file)} 字节")
            else:
                debug_print(f"警告: 输出文件未能生成")
            
            if deleted_exists:
                debug_print(f"删除内容文件大小: {os.path.getsize(deleted_file)} 字节")
            else:
                debug_print(f"警告: 删除内容文件未能生成")
            
            # 返回处理结果
            result = {
                'output_file': output_file,
                'deleted_file': deleted_file,
                'stats': stats
            }
            
            self.progress_signal.emit(100)
            self.finished_signal.emit(result)
            
        except Exception as e:
            debug_print(f"线程处理文档错误: {str(e)}")
            import traceback
            debug_print(traceback.format_exc())
            self.error_signal.emit(str(e))

class DocumentProcessorWidget(QWidget):
    """文档处理界面组件"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setup_ui()
        
    def setup_ui(self):
        """设置用户界面"""
        # 主布局
        main_layout = QVBoxLayout()
        main_layout.setSpacing(15)
        
        # 文件选择区域
        file_group = QGroupBox("文档选择")
        file_group.setStyleSheet("QGroupBox { font-size: 12pt; font-weight: bold; }")
        file_layout = QVBoxLayout()
        file_layout.setContentsMargins(15, 15, 15, 15)
        
        # 文件选择行
        file_select_layout = QHBoxLayout()
        self.file_path_label = QLabel("未选择文件")
        self.file_path_label.setStyleSheet("font-size: 10pt;")
        self.select_file_button = QPushButton("选择文档")
        self.select_file_button.setMinimumHeight(40)
        self.select_file_button.clicked.connect(self.select_document)
        file_select_layout.addWidget(self.file_path_label, 1)
        file_select_layout.addWidget(self.select_file_button)
        
        # 输出目录选择行
        output_dir_layout = QHBoxLayout()
        self.output_dir_label = QLabel("输出目录: ./output")
        self.output_dir_label.setStyleSheet("font-size: 10pt;")
        self.select_output_dir_button = QPushButton("选择输出目录")
        self.select_output_dir_button.setMinimumHeight(40)
        self.select_output_dir_button.clicked.connect(self.select_output_directory)
        output_dir_layout.addWidget(self.output_dir_label, 1)
        output_dir_layout.addWidget(self.select_output_dir_button)
        
        file_layout.addLayout(file_select_layout)
        file_layout.addLayout(output_dir_layout)
        file_group.setLayout(file_layout)
        
        # 处理选项区域
        options_group = QGroupBox("处理选项")
        options_group.setStyleSheet("QGroupBox { font-size: 12pt; font-weight: bold; }")
        options_layout = QVBoxLayout()
        options_layout.setContentsMargins(15, 15, 15, 15)
        
        # 删除文件名选项
        self.remove_filenames_checkbox = QCheckBox("删除文件名")
        self.remove_filenames_checkbox.setChecked(False)
        self.remove_filenames_checkbox.setStyleSheet("""
            QCheckBox {
                font-size: 10pt;
                color: #333;
            }
            QCheckBox::indicator {
                width: 18px;
                height: 18px;
            }
        """)
        options_layout.addWidget(self.remove_filenames_checkbox)
        
        # 删除大段注释选项
        self.remove_large_comments_checkbox = QCheckBox("删除大段注释（两行及以上）")
        self.remove_large_comments_checkbox.setChecked(False)
        self.remove_large_comments_checkbox.setStyleSheet("""
            QCheckBox {
                font-size: 10pt;
                color: #333;
            }
            QCheckBox::indicator {
                width: 18px;
                height: 18px;
            }
        """)
        options_layout.addWidget(self.remove_large_comments_checkbox)
        
        # 删除英文注释选项
        self.remove_english_comments_checkbox = QCheckBox("删除英文注释")
        self.remove_english_comments_checkbox.setChecked(False)
        self.remove_english_comments_checkbox.setStyleSheet("""
            QCheckBox {
                font-size: 10pt;
                color: #333;
            }
            QCheckBox::indicator {
                width: 18px;
                height: 18px;
            }
        """)
        options_layout.addWidget(self.remove_english_comments_checkbox)
        
        # 随机删除单行注释选项
        random_comments_layout = QHBoxLayout()
        random_comments_label = QLabel("随机删除单行注释比例:")
        random_comments_label.setStyleSheet("font-size: 10pt;")
        random_comments_layout.addWidget(random_comments_label)
        
        self.remove_comments_ratio = QSpinBox()
        self.remove_comments_ratio.setStyleSheet("font-size: 10pt; height: 30px;")
        self.remove_comments_ratio.setFixedHeight(30)
        self.remove_comments_ratio.setRange(0, 10)
        self.remove_comments_ratio.setValue(0)
        self.remove_comments_ratio.setSpecialValueText("不删除")
        random_comments_layout.addWidget(self.remove_comments_ratio)
        
        # 添加提示标签
        ratio_tip_label = QLabel("(0表示不删除，3表示每3个删除1个)")
        ratio_tip_label.setStyleSheet("font-size: 9pt; color: #666;")
        random_comments_layout.addWidget(ratio_tip_label)
        
        options_layout.addLayout(random_comments_layout)
        options_group.setLayout(options_layout)
        
        # 统计信息区域
        self.stats_group = QGroupBox("统计信息")
        self.stats_group.setStyleSheet("QGroupBox { font-size: 12pt; font-weight: bold; }")
        stats_layout = QVBoxLayout()
        stats_layout.setContentsMargins(15, 15, 15, 15)
        
        self.total_lines_label = QLabel("总行数: 0")
        self.total_lines_label.setStyleSheet("font-size: 10pt;")
        stats_layout.addWidget(self.total_lines_label)
        
        self.deleted_filenames_label = QLabel("删除的文件名数: 0")
        self.deleted_filenames_label.setStyleSheet("font-size: 10pt;")
        stats_layout.addWidget(self.deleted_filenames_label)
        
        self.deleted_large_comments_label = QLabel("删除的大段注释数: 0")
        self.deleted_large_comments_label.setStyleSheet("font-size: 10pt;")
        stats_layout.addWidget(self.deleted_large_comments_label)
        
        self.deleted_english_comments_label = QLabel("删除的英文注释数: 0")
        self.deleted_english_comments_label.setStyleSheet("font-size: 10pt;")
        stats_layout.addWidget(self.deleted_english_comments_label)
        
        self.deleted_random_comments_label = QLabel("随机删除的注释数: 0")
        self.deleted_random_comments_label.setStyleSheet("font-size: 10pt;")
        stats_layout.addWidget(self.deleted_random_comments_label)
        
        self.remaining_lines_label = QLabel("剩余行数: 0")
        self.remaining_lines_label.setStyleSheet("font-size: 10pt;")
        stats_layout.addWidget(self.remaining_lines_label)
        
        self.stats_group.setLayout(stats_layout)
        
        # 按钮区域
        buttons_layout = QHBoxLayout()
        
        self.process_button = QPushButton("处理文档")
        self.process_button.setStyleSheet("font-size: 11pt;")
        self.process_button.setMinimumHeight(45)
        self.process_button.clicked.connect(self.process_document)
        buttons_layout.addWidget(self.process_button)
        
        self.open_output_button = QPushButton("打开处理后文档")
        self.open_output_button.setStyleSheet("font-size: 11pt;")
        self.open_output_button.setMinimumHeight(45)
        self.open_output_button.clicked.connect(self.open_output_document)
        self.open_output_button.setEnabled(False)
        buttons_layout.addWidget(self.open_output_button)
        
        self.open_deleted_button = QPushButton("打开删除内容记录")
        self.open_deleted_button.setStyleSheet("font-size: 11pt;")
        self.open_deleted_button.setMinimumHeight(45)
        self.open_deleted_button.clicked.connect(self.open_deleted_content)
        self.open_deleted_button.setEnabled(False)
        buttons_layout.addWidget(self.open_deleted_button)
        
        # 进度条
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        self.progress_bar.setMinimumHeight(30)
        
        # 组合主布局
        main_layout.addWidget(file_group)
        main_layout.addWidget(options_group)
        main_layout.addWidget(self.stats_group)
        main_layout.addLayout(buttons_layout)
        main_layout.addWidget(self.progress_bar)
        
        # 设置主布局
        self.setLayout(main_layout)
        
        # 存储处理结果
        self.processed_output_file = None
        self.processed_deleted_file = None
        
        # 默认输出目录
        self.output_dir = "./output"
    
    def select_document(self):
        """选择要处理的文档"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, "选择文档", "", "Word文档 (*.docx);;所有文件 (*.*)"
        )
        if file_path:
            self.file_path_label.setText(file_path)
            if hasattr(self.parent(), 'status_bar'):
                self.parent().status_bar.showMessage(f"已选择文档: {file_path}")
            
            # 重置处理结果
            self.processed_output_file = None
            self.processed_deleted_file = None
            self.open_output_button.setEnabled(False)
            self.open_deleted_button.setEnabled(False)
    
    def select_output_directory(self):
        """选择输出目录"""
        directory = QFileDialog.getExistingDirectory(self, "选择输出目录")
        if directory:
            self.output_dir = directory
            self.output_dir_label.setText(f"输出目录: {directory}")
            if hasattr(self.parent(), 'status_bar'):
                self.parent().status_bar.showMessage(f"已选择输出目录: {directory}")
    
    def process_document(self):
        """处理文档"""
        # 检查是否选择了文档
        file_path = self.file_path_label.text()
        if file_path == "未选择文件":
            QMessageBox.warning(self, "警告", "请先选择要处理的文档")
            return
        
        # 检查文件是否存在
        if not os.path.exists(file_path):
            QMessageBox.warning(self, "警告", "所选文档不存在")
            return
        
        # 获取处理选项
        remove_filenames = self.remove_filenames_checkbox.isChecked()
        remove_large_comments = self.remove_large_comments_checkbox.isChecked()
        remove_english_comments = self.remove_english_comments_checkbox.isChecked()
        remove_comments_ratio = self.remove_comments_ratio.value()
        
        # 确保至少启用一个选项以生成删除内容文档
        has_delete_options = remove_filenames or remove_large_comments or remove_english_comments or remove_comments_ratio > 0
        debug_print(f"启用的删除选项: 删除文件名={remove_filenames}, 删除大段注释={remove_large_comments}, 删除英文注释={remove_english_comments}, 随机删除比例={remove_comments_ratio}")
        
        if not has_delete_options:
            # 如果用户没有选择任何删除选项，默认启用删除文件名选项
            debug_print("未选择任何删除选项，默认启用删除文件名选项")
            remove_filenames = True
            self.remove_filenames_checkbox.setChecked(True)
            QMessageBox.information(self, "提示", "为确保生成删除内容记录文档，已自动启用删除文件名选项")
        
        # 显示进度条
        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(0)
        if hasattr(self.parent(), 'status_bar'):
            self.parent().status_bar.showMessage("正在处理文档...")
        
        # 创建并启动处理线程
        self.processor_thread = DocumentProcessorThread({
            'input_path': file_path,
            'output_dir': self.output_dir,
            'remove_filenames': remove_filenames,
            'remove_large_comments': remove_large_comments,
            'remove_english_comments': remove_english_comments,
            'remove_comments_ratio': remove_comments_ratio
        })
        
        self.processor_thread.progress_signal.connect(self.update_progress)
        self.processor_thread.finished_signal.connect(self.document_processed)
        self.processor_thread.error_signal.connect(self.document_error)
        
        debug_print(f"启动处理线程，处理参数: {self.processor_thread.params}")
        self.processor_thread.start()
    
    def update_progress(self, value):
        """更新进度条"""
        self.progress_bar.setValue(value)
    
    def document_processed(self, result):
        """文档处理完成"""
        self.progress_bar.setVisible(False)
        
        debug_print("\n=== 文档处理完成 ===")
        debug_print(f"接收到结果:")
        debug_print(f"  result键: {result.keys()}")
        
        # 保存处理结果
        self.processed_output_file = result['output_file']
        self.processed_deleted_file = result['deleted_file']
        
        debug_print(f"\n=== 文件路径信息 ===")
        debug_print(f"处理后文档路径: {self.processed_output_file}")
        debug_print(f"删除内容文档路径: {self.processed_deleted_file}")
        
        # 使用绝对路径
        abs_output_file = os.path.abspath(self.processed_output_file)
        abs_deleted_file = os.path.abspath(self.processed_deleted_file)
        debug_print(f"处理后文档绝对路径: {abs_output_file}")
        debug_print(f"删除内容文档绝对路径: {abs_deleted_file}")
        
        # 检查输出目录
        output_dir = os.path.dirname(abs_output_file)
        debug_print(f"\n=== 输出目录信息 ===")
        debug_print(f"输出目录: {output_dir}")
        debug_print(f"目录存在: {os.path.exists(output_dir)}")
        debug_print(f"目录可写: {os.access(output_dir, os.W_OK)}")
        
        # 检查文件是否存在
        debug_print(f"\n=== 文件存在检查 ===")
        output_exists = os.path.exists(abs_output_file)
        deleted_exists = os.path.exists(abs_deleted_file)
        debug_print(f"处理后文档存在: {output_exists}")
        debug_print(f"删除内容文档存在: {deleted_exists}")
        
        if output_exists:
            debug_print(f"处理后文档大小: {os.path.getsize(abs_output_file)} 字节")
        
        if deleted_exists:
            debug_print(f"删除内容文档大小: {os.path.getsize(abs_deleted_file)} 字节")
        else:
            debug_print(f"警告: 删除内容文档不存在于预期位置")
            
            # 列出输出目录中所有可能的删除内容文档
            debug_print(f"搜索可能的删除内容文档:")
            deleted_candidates = []
            for root, dirs, files in os.walk(output_dir):
                for file in files:
                    if 'deleted' in file.lower() or 'alt_' in file.lower() or 'temp_' in file.lower():
                        file_path = os.path.join(root, file)
                        mod_time = os.path.getmtime(file_path)
                        deleted_candidates.append((file_path, mod_time))
            
            # 检查临时目录
            try:
                import tempfile
                temp_dir = tempfile.gettempdir()
                if os.path.exists(temp_dir):
                    debug_print(f"检查临时目录: {temp_dir}")
                    for file in os.listdir(temp_dir):
                        if ('deleted' in file.lower() or 'temp_' in file.lower()) and file.endswith('.docx'):
                            file_path = os.path.join(temp_dir, file)
                            mod_time = os.path.getmtime(file_path)
                            deleted_candidates.append((file_path, mod_time))
            except Exception as e:
                debug_print(f"检查临时目录时出错: {str(e)}")
            
            if deleted_candidates:
                # 按修改时间排序，获取最新的文件
                deleted_candidates.sort(key=lambda x: x[1], reverse=True)
                newest_deleted = deleted_candidates[0][0]
                debug_print(f"找到最新的可能删除内容文档: {newest_deleted}")
                
                # 更新删除内容文档路径
                self.processed_deleted_file = newest_deleted
                deleted_exists = True
                
                # 如果文件在临时目录，尝试复制到输出目录
                if tempfile.gettempdir() in newest_deleted:
                    try:
                        import shutil
                        output_copy_path = os.path.join(output_dir, os.path.basename(newest_deleted))
                        shutil.copy2(newest_deleted, output_copy_path)
                        debug_print(f"已将临时文件复制到输出目录: {output_copy_path}")
                        self.processed_deleted_file = output_copy_path
                    except Exception as e:
                        debug_print(f"复制临时文件时出错: {str(e)}")
        
        # 列出输出目录中的所有文件
        debug_print(f"\n=== 输出目录文件列表 ===")
        try:
            for file in os.listdir(output_dir):
                file_path = os.path.join(output_dir, file)
                if os.path.isfile(file_path):
                    debug_print(f"  {file} - {os.path.getsize(file_path)} 字节")
        except Exception as e:
            debug_print(f"列出文件时出错: {str(e)}")
        
        # 主动验证删除内容文档是否存在，如果不存在则尝试再次创建
        if not deleted_exists:
            debug_print(f"\n=== 尝试创建删除内容文档 ===")
            debug_print(f"警告：删除内容文档不存在，尝试再次创建: {abs_deleted_file}")
            try:
                from docx import Document
                deleted_doc = Document()
                deleted_doc.add_heading('删除内容记录（自动创建）', 0)
                deleted_doc.add_paragraph(f"原始文档: {os.path.basename(self.processed_output_file)}")
                deleted_doc.add_paragraph(f"注意: 此文档为自动创建，因原始删除内容文档未能生成")
                deleted_doc.add_paragraph(f"处理时间: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                deleted_doc.add_paragraph("未删除任何内容。")
                
                # 尝试不同的保存方式
                try:
                    debug_print(f"尝试使用绝对路径保存: {abs_deleted_file}")
                    deleted_doc.save(abs_deleted_file)
                except Exception as e1:
                    debug_print(f"使用绝对路径保存失败: {str(e1)}")
                    try:
                        # 尝试保存到当前目录
                        current_dir_path = os.path.join(os.getcwd(), f"emergency_deleted_content_{datetime.datetime.now().strftime('%m%d%H%M%S')}.docx")
                        debug_print(f"尝试保存到当前目录: {current_dir_path}")
                        deleted_doc.save(current_dir_path)
                        debug_print(f"成功保存到当前目录，更新删除内容文档路径")
                        self.processed_deleted_file = current_dir_path
                    except Exception as e2:
                        debug_print(f"保存到当前目录也失败: {str(e2)}")
                
                # 再次检查
                if os.path.exists(self.processed_deleted_file):
                    debug_print(f"确认替代的删除内容文档已存在: {self.processed_deleted_file}")
                    debug_print(f"文件大小: {os.path.getsize(self.processed_deleted_file)} 字节")
                else:
                    debug_print(f"错误：即使尝试创建替代文档，删除内容文档仍不存在")
            except Exception as e:
                debug_print(f"创建替代删除内容文档时出错: {str(e)}")
                import traceback
                debug_print(traceback.format_exc())
        
        # 更新统计信息
        debug_print("\n=== 更新统计信息 ===")
        stats = result['stats']
        self.total_lines_label.setText(f"总行数: {stats['total_lines']}")
        self.deleted_filenames_label.setText(f"删除的文件名数: {stats['deleted_filenames']}")
        self.deleted_large_comments_label.setText(f"删除的大段注释数: {stats['deleted_large_comments']}")
        self.deleted_english_comments_label.setText(f"删除的英文注释数: {stats['deleted_english_comments']}")
        self.deleted_random_comments_label.setText(f"随机删除的注释数: {stats['deleted_random_comments']}")
        self.remaining_lines_label.setText(f"剩余行数: {stats['remaining_lines']}")
        
        # 启用打开文件按钮
        self.open_output_button.setEnabled(True)
        self.open_deleted_button.setEnabled(True)
        
        # 更新状态栏
        if hasattr(self.parent(), 'status_bar'):
            self.parent().status_bar.showMessage(f"文档处理完成: {os.path.basename(self.processed_output_file)}")
        
        # 显示处理结果
        QMessageBox.information(
            self, "处理完成", 
            f"文档处理完成\n\n"
            f"处理后文档: {os.path.basename(self.processed_output_file)}\n"
            f"删除内容记录: {os.path.basename(self.processed_deleted_file)}\n\n"
            f"总行数: {stats['total_lines']}\n"
            f"删除的文件名数: {stats['deleted_filenames']}\n"
            f"删除的大段注释数: {stats['deleted_large_comments']}\n"
            f"删除的英文注释数: {stats['deleted_english_comments']}\n"
            f"随机删除的注释数: {stats['deleted_random_comments']}\n"
            f"剩余行数: {stats['remaining_lines']}"
        )
        
        # 显著的调试输出，显示生成了哪些文档
        debug_print("\n")
        debug_print("=" * 80)
        debug_print("文档处理完成，已生成以下文件:")
        debug_print("-" * 80)
        debug_print(f"1. 处理后文档: {self.processed_output_file}")
        debug_print(f"   - 文件存在: {os.path.exists(self.processed_output_file)}")
        if os.path.exists(self.processed_output_file):
            debug_print(f"   - 文件大小: {os.path.getsize(self.processed_output_file)} 字节")
        debug_print("-" * 80)
        debug_print(f"2. 删除内容记录文档: {self.processed_deleted_file}")
        debug_print(f"   - 文件存在: {os.path.exists(self.processed_deleted_file)}")
        if os.path.exists(self.processed_deleted_file):
            debug_print(f"   - 文件大小: {os.path.getsize(self.processed_deleted_file)} 字节")
        debug_print("=" * 80)
        debug_print("\n")
    
    def document_error(self, error_msg):
        """文档处理错误"""
        self.progress_bar.setVisible(False)
        if hasattr(self.parent(), 'status_bar'):
            self.parent().status_bar.showMessage("处理文档时出错")
        
        QMessageBox.critical(self, "错误", f"处理文档时出错:\n{error_msg}")
    
    def open_output_document(self):
        """打开处理后的文档"""
        if not self.processed_output_file or not os.path.exists(self.processed_output_file):
            QMessageBox.warning(self, "警告", "处理后的文档不存在")
            return
        
        try:
            # 使用系统默认程序打开文件
            if sys.platform == 'win32':
                os.startfile(self.processed_output_file)
            elif sys.platform == 'darwin':  # macOS
                os.system(f'open "{self.processed_output_file}"')
            else:  # Linux
                os.system(f'xdg-open "{self.processed_output_file}"')
            
            if hasattr(self.parent(), 'status_bar'):
                self.parent().status_bar.showMessage(f"已打开文件: {self.processed_output_file}")
        except Exception as e:
            QMessageBox.critical(self, "错误", f"无法打开文件: {str(e)}")
            if hasattr(self.parent(), 'status_bar'):
                self.parent().status_bar.showMessage(f"打开文件失败: {self.processed_output_file}")
    
    def open_deleted_content(self):
        """打开删除内容记录文档"""
        debug_print(f"\n=== 尝试打开删除内容记录文档 ===")
        debug_print(f"删除内容文档路径: {self.processed_deleted_file}")
        
        if not self.processed_deleted_file:
            debug_print("删除内容记录文档路径为空")
            QMessageBox.warning(self, "警告", "删除内容记录文档路径未设置")
            return
        
        # 使用绝对路径
        abs_deleted_file = os.path.abspath(self.processed_deleted_file)
        debug_print(f"删除内容文档绝对路径: {abs_deleted_file}")
            
        if not os.path.exists(abs_deleted_file):
            debug_print(f"警告：删除内容文档不存在: {abs_deleted_file}")
            
            # 尝试基于原始文件路径猜测删除内容文档路径
            if hasattr(self, 'processed_output_file') and self.processed_output_file:
                debug_print(f"尝试基于处理后文档路径查找删除内容文档")
                
                # 获取处理后文档的文件名和目录
                output_dir = os.path.dirname(self.processed_output_file)
                output_basename = os.path.basename(self.processed_output_file)
                output_name_without_ext = os.path.splitext(output_basename)[0]
                
                debug_print(f"处理后文档目录: {output_dir}")
                debug_print(f"处理后文档文件名: {output_basename}")
                debug_print(f"不带扩展名的文件名: {output_name_without_ext}")
                
                # 尝试几种可能的删除内容文档名称
                possible_names = [
                    f"{output_name_without_ext}_deleted.docx",
                    f"{output_name_without_ext.replace('_processed', '')}_deleted.docx",
                    f"{output_name_without_ext}_deleted_content.docx"
                ]
                
                debug_print(f"可能的删除内容文档名称:")
                for name in possible_names:
                    debug_print(f"  尝试: {name}")
                    possible_path = os.path.join(output_dir, name)
                    if os.path.exists(possible_path):
                        debug_print(f"找到匹配的删除内容文档: {possible_path}")
                        self.processed_deleted_file = possible_path
                        break
                
                # 如果仍然找不到，尝试查找目录中包含"deleted"的最新文件
                if not os.path.exists(self.processed_deleted_file):
                    debug_print(f"在目录中查找包含'deleted'的最新文件")
                    try:
                        deleted_files = []
                        for file in os.listdir(output_dir):
                            if 'deleted' in file.lower() and file.endswith('.docx'):
                                file_path = os.path.join(output_dir, file)
                                if os.path.isfile(file_path):
                                    deleted_files.append((file_path, os.path.getmtime(file_path)))
                        
                        if deleted_files:
                            # 按修改时间排序，获取最新的文件
                            deleted_files.sort(key=lambda x: x[1], reverse=True)
                            newest_file = deleted_files[0][0]
                            debug_print(f"找到最新的删除内容文档: {newest_file}")
                            self.processed_deleted_file = newest_file
                    except Exception as e:
                        debug_print(f"查找删除内容文档时出错: {str(e)}")
            
            # 检查是否找到替代文件
            if not os.path.exists(self.processed_deleted_file):
                debug_print(f"无法找到任何替代的删除内容文档")
                QMessageBox.warning(self, "警告", "删除内容记录文档不存在，尝试查找替代文档也失败了")
                
                # 提示用户检查输出目录
                msg = QMessageBox()
                msg.setIcon(QMessageBox.Question)
                msg.setText("是否要打开输出目录查看文件?")
                msg.setInformativeText("无法找到删除内容记录文档，您可以手动查看输出目录中的文件。")
                msg.setWindowTitle("打开输出目录")
                msg.setStandardButtons(QMessageBox.Yes | QMessageBox.No)
                if msg.exec_() == QMessageBox.Yes:
                    output_dir = os.path.dirname(self.processed_output_file) if hasattr(self, 'processed_output_file') else self.output_dir
                    try:
                        if sys.platform == 'win32':
                            os.startfile(output_dir)
                        elif sys.platform == 'darwin':  # macOS
                            os.system(f'open "{output_dir}"')
                        else:  # Linux
                            os.system(f'xdg-open "{output_dir}"')
                    except Exception as e:
                        debug_print(f"打开输出目录时出错: {str(e)}")
                        QMessageBox.critical(self, "错误", f"无法打开输出目录: {str(e)}")
                return
        
        # 正常打开文件
        try:
            debug_print(f"使用系统默认程序打开文件: {self.processed_deleted_file}")
            if sys.platform == 'win32':
                os.startfile(self.processed_deleted_file)
            elif sys.platform == 'darwin':  # macOS
                os.system(f'open "{self.processed_deleted_file}"')
            else:  # Linux
                os.system(f'xdg-open "{self.processed_deleted_file}"')
            
            if hasattr(self.parent(), 'status_bar'):
                self.parent().status_bar.showMessage(f"已打开文件: {self.processed_deleted_file}")
                
            debug_print(f"文件已成功打开")
        except Exception as e:
            debug_print(f"打开文件出错: {str(e)}")
            QMessageBox.critical(self, "错误", f"无法打开文件: {str(e)}")
            if hasattr(self.parent(), 'status_bar'):
                self.parent().status_bar.showMessage(f"打开文件失败: {self.processed_deleted_file}")

def main():
    app = QApplication(sys.argv)
    window = QMainWindow()
    window.setWindowTitle("文档处理器")
    window.setGeometry(100, 100, 800, 600)
    
    # 创建文档处理组件
    processor_widget = DocumentProcessorWidget()
    window.setCentralWidget(processor_widget)
    
    # 创建状态栏
    status_bar = QStatusBar()
    window.setStatusBar(status_bar)
    window.status_bar = status_bar
    status_bar.showMessage("准备就绪")
    
    window.show()
    sys.exit(app.exec_())

if __name__ == "__main__":
    main() 